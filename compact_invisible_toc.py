import re
import sys
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def set_font(run, size=13, bold=False):
    run.font.name = 'TH Sarabun New'
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement('w:rFonts')
        rpr.append(fonts)
    for name in ('ascii', 'hAnsi', 'cs'):
        fonts.set(qn(f'w:{name}'), 'TH Sarabun New')
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def clear_paragraph(paragraph):
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1


def add_pageref(paragraph, bookmark):
    run = paragraph.add_run()
    set_font(run)
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instruction = OxmlElement('w:instrText')
    instruction.set(qn('xml:space'), 'preserve')
    instruction.text = f' PAGEREF {bookmark} \\h '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    cached = OxmlElement('w:t')
    cached.text = '?'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend((begin, instruction, separate, cached, end))


def set_no_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        node = borders.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            borders.append(node)
        node.set(qn('w:val'), 'nil')
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            cell_borders = tc_pr.find(qn('w:tcBorders'))
            if cell_borders is None:
                cell_borders = OxmlElement('w:tcBorders')
                tc_pr.append(cell_borders)
            for edge in ('top', 'left', 'bottom', 'right'):
                node = cell_borders.find(qn(f'w:{edge}'))
                if node is None:
                    node = OxmlElement(f'w:{edge}')
                    cell_borders.append(node)
                node.set(qn('w:val'), 'nil')


def set_geometry(table, total_width, title_width, page_width):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(total_width))
    tbl_w.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, (title_width, page_width)):
        grid_col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, (title_width, page_width)):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')
            tc_mar = tc_pr.find(qn('w:tcMar'))
            if tc_mar is None:
                tc_mar = OxmlElement('w:tcMar')
                tc_pr.append(tc_mar)
            for side in ('top', 'start', 'bottom', 'end'):
                node = tc_mar.find(qn(f'w:{side}'))
                if node is None:
                    node = OxmlElement(f'w:{side}')
                    tc_mar.append(node)
                node.set(qn('w:w'), '0')
                node.set(qn('w:type'), 'dxa')


def bookmark_name(paragraph):
    for element in paragraph._p.iter(qn('w:bookmarkStart')):
        name = element.get(qn('w:name'))
        if name and name.startswith('toc_entry_'):
            return name
    return None


source, output = sys.argv[1], sys.argv[2]
document = Document(source)

# Keep only chapter-level headings and meaningful numbered section headings.
entries = []
for paragraph in document.paragraphs:
    text = paragraph.text.strip()
    style = paragraph.style.name
    if style == 'Heading 1' and text != '3. สารบัญ':
        entries.append(paragraph)
    elif style == 'Heading 2' and (
        re.match(r'^(5\.[1-6]|10\.[12])\s+', text) or text == 'ภาคผนวก ก'
    ):
        entries.append(paragraph)

missing = [p.text.strip() for p in entries if bookmark_name(p) is None]
if missing:
    raise RuntimeError(f'ไม่พบ bookmark สำหรับหัวข้อ: {missing}')

old_toc = next(
    table for table in document.tables
    if len(table.columns) == 2 and table.cell(0, 0).text.strip() == 'หัวข้อ'
)
old_toc._tbl.getparent().remove(old_toc._tbl)

main_heading = next(p for p in document.paragraphs if p.text.strip().startswith('4. วัตถุประสงค์และเป้าหมาย'))

# Keep the section break after the replacement TOC.  This preserves the
# front-matter page numbering for the TOC and lets the main text start at 1.
section_break = None
cursor = main_heading._p.getprevious()
while cursor is not None:
    p_pr = cursor.find(qn('w:pPr'))
    if p_pr is not None and p_pr.find(qn('w:sectPr')) is not None:
        section_break = cursor
        break
    cursor = cursor.getprevious()
if section_break is None:
    raise RuntimeError('ไม่พบตัวแบ่งส่วนหลังสารบัญ')

table = document.add_table(rows=len(entries), cols=2)
table.style = 'Table Grid'
# XML table widths use twips (DXA), whereas python-docx page dimensions are
# stored internally in EMU.  Using twips keeps the page-number column within
# the right margin instead of extending beyond the paper.
section = document.sections[0]
total_width = int(section.page_width.twips - section.left_margin.twips - section.right_margin.twips)
page_width = 540
set_geometry(table, total_width, total_width - page_width, page_width)
set_no_borders(table)

for row, heading in zip(table.rows, entries):
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    title = row.cells[0].paragraphs[0]
    clear_paragraph(title)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if heading.style.name == 'Heading 2':
        title.paragraph_format.left_indent = Pt(15)
    run = title.add_run(heading.text.strip())
    set_font(run, size=13, bold=(heading.style.name == 'Heading 1'))
    page = row.cells[1].paragraphs[0]
    clear_paragraph(page)
    page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_pageref(page, bookmark_name(heading))

section_break.addprevious(table._tbl)
document.save(output)
print(f'Replaced the long TOC with {len(entries)} concise invisible-table entries.')
