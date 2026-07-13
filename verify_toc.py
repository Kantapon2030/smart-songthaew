from docx import Document
from zipfile import ZipFile
import sys

path = sys.argv[1]
with ZipFile(path) as archive:
    assert archive.testzip() is None
    document_xml = archive.read('word/document.xml').decode('utf-8')
    settings_xml = archive.read('word/settings.xml').decode('utf-8')

document = Document(path)
all_text = '\n'.join(paragraph.text for paragraph in document.paragraphs)
assert '[[TOC]]' not in all_text
assert 'TOC \\o "1-3"' in document_xml
assert 'updateFields' in settings_xml
for style_name in ('Heading 1', 'Heading 2', 'Heading 3'):
    assert any(paragraph.style.name == style_name for paragraph in document.paragraphs)

print('PASS: TOC field, update-on-open setting, three heading levels, placeholder removal, and DOCX integrity verified.')
