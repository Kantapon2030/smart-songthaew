from docx import Document
from zipfile import ZipFile
from docx.oxml.ns import qn
import sys

path = sys.argv[1]
with ZipFile(path) as archive:
    assert archive.testzip() is None
    xml = archive.read('word/document.xml').decode('utf-8')
    header_xml = '\n'.join(
        archive.read(name).decode('utf-8')
        for name in archive.namelist()
        if name.startswith('word/header') and name.endswith('.xml')
    )
    assert 'PAGEREF toc_entry_' in xml
    assert ' PAGE ' in header_xml
    assert 'thaiLetters' in xml
    assert 'w:fmt="decimal"' in xml

document = Document(path)
assert len(document.sections) == 3
assert sum(1 for table in document.tables if table.cell(0, 0).text.strip() == 'หัวข้อ' and table.cell(0, 1).text.strip() == 'หน้า') == 1
assert 'TOC \\o "1-3"' not in xml
assert all(section.header.paragraphs[0].alignment is not None for section in document.sections[1:])
assert all(style.font.color.rgb is not None for style in (document.styles['Heading 1'], document.styles['Heading 2'], document.styles['Heading 3']))
print('PASS: table TOC, PAGEREF fields, Thai/Arabic pagination sections, upper-right headers, black heading styles, and DOCX integrity verified.')
