from copy import deepcopy
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import sys


def ensure_ppr(paragraph_element):
    ppr = paragraph_element.find(qn('w:pPr'))
    if ppr is None:
        ppr = OxmlElement('w:pPr')
        paragraph_element.insert(0, ppr)
    return ppr


def set_page_number_format(sect_pr, fmt, start=None):
    node = sect_pr.find(qn('w:pgNumType'))
    if node is None:
        node = OxmlElement('w:pgNumType')
        sect_pr.append(node)
    node.set(qn('w:fmt'), fmt)
    if start is not None:
        node.set(qn('w:start'), str(start))
    elif qn('w:start') in node.attrib:
        del node.attrib[qn('w:start')]


def set_run_font(run, size=14, bold=False):
    run.font.name = 'TH Sarabun New'
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('ascii', 'hAnsi', 'cs'):
        rfonts.set(qn(f'w:{attr}'), 'TH Sarabun New')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def clear_paragraph(paragraph):
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def append_field(paragraph, instruction, cached='?'):
    run = paragraph.add_run()
    set_run_font(run)
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' {instruction} '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    value = OxmlElement('w:t')
    value.text = cached
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend((begin, instr, separate, value, end))


def set_header_page_number(section, show_number):
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for paragraph in section.footer.paragraphs:
        clear_paragraph(paragraph)
    for paragraph in section.header.paragraphs:
        clear_paragraph(paragraph)
    if show_number:
        paragraph = section.header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        append_field(paragraph, 'PAGE')


def set_cell_margins(cell, top=70, start=100, bottom=70, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn('w:tcMar'))
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def shade_cell(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn('w:shd'))
    if shading is None:
        shading = OxmlElement('w:shd')
        tc_pr.append(shading)
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:fill'), color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement('w:tblHeader')
    flag.set(qn('w:val'), 'true')
    tr_pr.append(flag)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement('w:cantSplit'))


def set_table_geometry(table, total_width, first_width, second_width):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(total_width))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '0')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, (first_width, second_width)):
        grid_col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, (first_width, second_width)):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')


source, output = sys.argv[1], sys.argv[2]
document = Document(source)

# Restore the original formal black heading treatment.
for style_name in ('Heading 1', 'Heading 2', 'Heading 3', 'Heading 4'):
    style = document.styles[style_name]
    style.font.name = 'TH Sarabun New'
    style.font.color.rgb = RGBColor(0, 0, 0)

paragraphs = document.paragraphs
ack = next(p for p in paragraphs if p.text.strip() == 'กิตติกรรมประกาศ')
toc_heading = next(p for p in paragraphs if p.text.strip() == '3. สารบัญ')
main_heading = next(p for p in paragraphs if p.text.strip().startswith('4. วัตถุประสงค์และเป้าหมาย'))

# Capture TOC entries before changing the document body. Heading 4 is intentionally
# omitted to keep the printed TOC concise and limited to three formal levels.
entries = [
    p for p in paragraphs
    if p.style.name in ('Heading 1', 'Heading 2', 'Heading 3')
    and p.text.strip() != '3. สารบัญ'
]

# Clear the previous Word TOC field/placeholder between its title and the first
# main-content heading. The replacement is a true table rather than dot leaders.
body = document._body._body
cursor = toc_heading._p.getnext()
while cursor is not None and cursor is not main_heading._p:
    next_cursor = cursor.getnext()
    body.remove(cursor)
    cursor = next_cursor

# Give every TOC entry a stable bookmark for its dynamic PAGEREF page-number cell.
existing_ids = []
for element in document.element.body.iter(qn('w:bookmarkStart')):
    value = element.get(qn('w:id'))
    if value and value.isdigit():
        existing_ids.append(int(value))
next_bookmark_id = max(existing_ids, default=0) + 1
bookmarks = []
for index, paragraph in enumerate(entries, start=1):
    name = f'toc_entry_{index:03d}'
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(next_bookmark_id))
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(next_bookmark_id))
    paragraph._p.insert(1, start)
    paragraph._p.append(end)
    bookmarks.append((paragraph, name))
    next_bookmark_id += 1

table = document.add_table(rows=len(bookmarks) + 1, cols=2)
table.style = 'Table Grid'
table.autofit = False
section_width = document.sections[0].page_width - document.sections[0].left_margin - document.sections[0].right_margin
total_width = int(section_width)
page_width = 720
title_width = total_width - page_width
set_table_geometry(table, total_width, title_width, page_width)

header = table.rows[0]
set_repeat_header(header)
for column, label in enumerate(('หัวข้อ', 'หน้า')):
    cell = header.cells[column]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    shade_cell(cell, 'D9E2F3')
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(label)
    set_run_font(run, size=15, bold=True)

for row_index, (heading, bookmark_name) in enumerate(bookmarks, start=1):
    row = table.rows[row_index]
    prevent_row_split(row)
    title_cell, page_cell = row.cells
    for cell in (title_cell, page_cell):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
    title_para = title_cell.paragraphs[0]
    clear_paragraph(title_para)
    level = int(heading.style.name[-1])
    title_para.paragraph_format.left_indent = Pt((level - 1) * 12)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run(heading.text.strip())
    set_run_font(run, size=14, bold=(level == 1))
    page_para = page_cell.paragraphs[0]
    clear_paragraph(page_para)
    page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    append_field(page_para, f'PAGEREF {bookmark_name} \\h')

# Place the table directly after the TOC title.
main_heading._p.addprevious(table._tbl)

# Create section boundaries: cover (no number), preliminary pages (Thai letters),
# and main content (Arabic numbers restarting at 1).
section_prototype = deepcopy(document._body._body.sectPr)
cover_end = paragraphs[paragraphs.index(ack) - 1]
cover_pr = ensure_ppr(cover_end._p)
cover_pr.append(deepcopy(section_prototype))

section_break = document.add_paragraph()
section_break._p.addprevious(deepcopy(section_break._p)) if False else None
main_heading._p.addprevious(section_break._p)
front_pr = ensure_ppr(section_break._p)
front_sect = deepcopy(section_prototype)
set_page_number_format(front_sect, 'thaiLetters', start=1)
front_pr.append(front_sect)
set_page_number_format(document._body._body.sectPr, 'decimal', start=1)

# Re-query the section list after inserting the two boundaries.
sections = document.sections
if len(sections) != 3:
    raise RuntimeError(f'Expected 3 sections, found {len(sections)}')
document.settings.odd_and_even_pages_header_footer = False
set_header_page_number(sections[0], show_number=False)
set_header_page_number(sections[1], show_number=True)
set_header_page_number(sections[2], show_number=True)

settings = document.settings.element
update = settings.find(qn('w:updateFields'))
if update is None:
    update = OxmlElement('w:updateFields')
    settings.append(update)
update.set(qn('w:val'), 'true')

document.save(output)
print(f'Created a {len(bookmarks)}-entry table of contents and configured {len(sections)} pagination sections.')
