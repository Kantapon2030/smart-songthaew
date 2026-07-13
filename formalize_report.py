from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import sys


def set_heading_text(paragraph, text, style_name):
    paragraph.clear()
    paragraph.style = style_name
    run = paragraph.add_run(text)
    run.font.name = 'TH Sarabun New'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'TH Sarabun New')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'TH Sarabun New')
    run._element.rPr.rFonts.set(qn('w:cs'), 'TH Sarabun New')
    run.font.size = Pt(16)
    run.bold = True


def add_page_field(paragraph):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    run = paragraph.add_run()
    run.font.name = 'TH Sarabun New'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'TH Sarabun New')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'TH Sarabun New')
    run._element.rPr.rFonts.set(qn('w:cs'), 'TH Sarabun New')
    run.font.size = Pt(14)
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instruction = OxmlElement('w:instrText')
    instruction.set(qn('xml:space'), 'preserve')
    instruction.text = ' PAGE '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    cached = OxmlElement('w:t')
    cached.text = '1'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend((begin, instruction, separate, cached, end))


source, output = sys.argv[1], sys.argv[2]
document = Document(source)

# Correct the two internal numbered groups that were being incorrectly promoted
# to top-level entries in the table of contents.
replacements = {
    '1. ฝั่งอุปกรณ์ติดตั้งบนรถสองแถว': ('5.1.1 ฝั่งอุปกรณ์ติดตั้งบนรถสองแถว', 'Heading 3'),
    '2. ฝั่งระบบสื่อสาร VIBE': ('5.1.2 ฝั่งระบบสื่อสาร VIBE', 'Heading 3'),
    '3. ฝั่งสถานีรับข้อมูลและระบบกลาง': ('5.1.3 ฝั่งสถานีรับข้อมูลและระบบกลาง', 'Heading 3'),
    '4. ฝั่งเว็บแอปพลิเคชันสำหรับผู้โดยสาร': ('5.1.4 ฝั่งเว็บแอปพลิเคชันสำหรับผู้โดยสาร', 'Heading 3'),
    '1. ฟังก์ชันของอุปกรณ์ VIBE บนรถ': ('5.4.4.1 ฟังก์ชันของอุปกรณ์ VIBE บนรถ', 'Heading 4'),
    '2. ฟังก์ชันของสถานีฐาน': ('5.4.4.2 ฟังก์ชันของสถานีฐาน', 'Heading 4'),
    '3. ฟังก์ชันของระบบกลาง': ('5.4.4.3 ฟังก์ชันของระบบกลาง', 'Heading 4'),
    '4. ฟังก์ชันของเว็บแอปพลิเคชัน': ('5.4.4.4 ฟังก์ชันของเว็บแอปพลิเคชัน', 'Heading 4'),
}

found = set()
for paragraph in document.paragraphs:
    text = paragraph.text.strip()
    if text in replacements:
        new_text, style_name = replacements[text]
        set_heading_text(paragraph, new_text, style_name)
        found.add(text)

missing = set(replacements) - found
if missing:
    raise RuntimeError(f'ไม่พบหัวข้อที่ต้องจัดระเบียบ: {sorted(missing)}')

# Put a centered PAGE field in the footer of every section, including the first page.
document.settings.odd_and_even_pages_header_footer = False
for section in document.sections:
    section.different_first_page_header_footer = False
    section.footer.is_linked_to_previous = False
    add_page_field(section.footer.paragraphs[0])

# Ask Word to update TOC/PAGE fields when the file is opened.
settings = document.settings.element
update = settings.find(qn('w:updateFields'))
if update is None:
    update = OxmlElement('w:updateFields')
    settings.append(update)
update.set(qn('w:val'), 'true')

document.save(output)
print(f'Formalized {len(found)} internal headings and added PAGE fields to {len(document.sections)} section(s).')
