from docx import Document
import sys

doc = Document(sys.argv[1])
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip().replace("\t", " ")
    if "5.2.7" in t or "5.2.12" in t or "สรุปเทคโนโลยี" in t:
        print(f"P {i}: style={p.style.name!r} text={t!r}")

for ti, table in enumerate(doc.tables):
    text = " | ".join(cell.text.replace("\n", " / ") for row in table.rows for cell in row.cells)
    if "Power Management" in text or "5.2.12" in text:
        print(f"TABLE {ti}: rows={len(table.rows)} cols={len(table.columns)}")
        for ri, row in enumerate(table.rows):
            print(f"  R {ri}: " + " || ".join(cell.text.replace("\n", " / ") for cell in row.cells))

print("--- CONTEXT 5.2.7 ---")
for i in range(295, min(365, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    print(f"P {i}: style={p.style.name!r} align={p.alignment} text={p.text!r}")

print("--- FORMAT SAMPLES ---")
for i in (310, 311, 340, 343):
    p = doc.paragraphs[i]
    print(f"P {i} pPr={p._p.pPr.xml if p._p.pPr is not None else None}")
    for ri, r in enumerate(p.runs):
        print(f"  run {ri}: text={r.text!r} rPr={r._r.rPr.xml if r._r.rPr is not None else None}")
for ti in (3,):
    table = doc.tables[ti]
    print(f"T {ti} style={table.style.name!r} tblPr={table._tbl.tblPr.xml}")
    for ri in (0, 1, 6):
        cell = table.rows[ri].cells[0]
        print(f" T{ti}R{ri}C0 tcPr={cell._tc.tcPr.xml if cell._tc.tcPr is not None else None}")
