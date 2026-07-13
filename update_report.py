from copy import deepcopy
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

SOURCE = r"รายงานฉบับสมบูรณ์ รหัสโครงการ 28P23S00194_ปรับปรุง.docx"


def copy_ppr(target, source):
    if source._p.pPr is not None:
        target._p.insert(0, deepcopy(source._p.pPr))


def copy_rpr(target_run, source_run):
    if source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))


def remove_all_runs(paragraph):
    for run in paragraph.runs:
        paragraph._p.remove(run._r)


def insert_before(anchor, element):
    anchor._p.addprevious(element)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)
    shd.set(qn('w:val'), 'clear')


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    tr_pr.append(cant_split)


def set_cell_text(cell, text, template_run, align, bold=False):
    p = cell.paragraphs[0]
    remove_all_runs(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    copy_rpr(run, template_run)
    run.bold = bold
    run.font.size = Pt(14)
    return p


def add_formula_paragraph(doc, anchor, parts, body_template):
    p = doc.add_paragraph()
    copy_ppr(p, body_template)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for text, subscript in parts:
        run = p.add_run(text)
        copy_rpr(run, body_template.runs[0])
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(15)
        run.font.subscript = subscript
    insert_before(anchor, p._p)


def add_body(doc, anchor, text, body_template):
    p = doc.add_paragraph()
    copy_ppr(p, body_template)
    p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    run = p.add_run(text)
    copy_rpr(run, body_template.runs[0])
    insert_before(anchor, p._p)


def add_heading(doc, anchor, text, heading_template):
    p = doc.add_paragraph()
    copy_ppr(p, heading_template)
    p.alignment = None
    run = p.add_run(text)
    copy_rpr(run, heading_template.runs[0])
    insert_before(anchor, p._p)


def add_caption(doc, anchor, text, caption_template):
    p = doc.add_paragraph()
    copy_ppr(p, caption_template)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    copy_rpr(run, caption_template.runs[0])
    insert_before(anchor, p._p)


doc = Document(SOURCE)
paragraphs = doc.paragraphs
heading_template = next(p for p in paragraphs if p.text.strip() == '5.2.7 การจัดการพลังงานแบบปรับตัว')
body_template = paragraphs[311]
caption_template = paragraphs[340]
anchor = next(p for p in paragraphs if p.text.strip() == '5.2.8 ระบบเก็บข้อมูลชั่วคราวและส่งต่อภายหลัง')

add_heading(doc, anchor, '5.2.7.1 การประเมินระยะเวลาการใช้งานและการชาร์จพลังงานแสงอาทิตย์เชิงทฤษฎี', heading_template)
add_body(doc, anchor, 'นอกจากการจัดการพลังงานแบบปรับตัวแล้ว อุปกรณ์ VIBE ยังสามารถประเมินประสิทธิภาพด้านพลังงานในเชิงทฤษฎีได้จากสเปกของอุปกรณ์ที่ใช้จริง ได้แก่ แบตเตอรี่ แผงโซลาร์เซลล์ และกำลังไฟฟ้าโดยประมาณของอุปกรณ์ IoT ภายในระบบ การประเมินนี้ช่วยให้สามารถคาดการณ์ระยะเวลาการใช้งานสูงสุดของอุปกรณ์ และเวลาที่ต้องใช้ในการชาร์จแบตเตอรี่ด้วยพลังงานแสงอาทิตย์ก่อนนำไปเปรียบเทียบกับผลการทดสอบจริงภาคสนาม', body_template)
add_body(doc, anchor, 'อุปกรณ์ VIBE ในระบบต้นแบบใช้แบตเตอรี่ Li-ion ขนาด 3.7V 5200mAh ร่วมกับแผงโซลาร์เซลล์ขนาด 6V 550mA โดยอุปกรณ์หลักที่ใช้พลังงานประกอบด้วยบอร์ด ESP8266 โมดูล GPS และโมดูลสื่อสาร LoRa SX1276 ซึ่งทำงานร่วมกันในการรับพิกัด จัดรูปแบบข้อมูล และส่งข้อมูลผ่านระบบ VIBE นอกจากนี้ ระบบยังมีแนวคิดปรับความถี่การส่งข้อมูลตามสถานะการเคลื่อนที่ของรถและระดับพลังงาน เพื่อช่วยลดการใช้พลังงานและยืดระยะเวลาการทำงานของอุปกรณ์', body_template)
add_caption(doc, anchor, 'ตารางที่ 5.2.7.1 การประเมินเชิงทฤษฎีของระบบพลังงานอุปกรณ์ VIBE', caption_template)

rows = [
    ('รายการ', 'ค่าที่ใช้ในการประเมินเชิงทฤษฎี', 'หมายเหตุ'),
    ('ความจุแบตเตอรี่', '5,200 mAh', 'ค่าตามสเปกแบตเตอรี่ที่ใช้ในต้นแบบ'),
    ('แรงดันแบตเตอรี่', '3.7 V', 'แบตเตอรี่ Li-ion 1 cell'),
    ('พลังงานรวมของแบตเตอรี่', '19.24 Wh', '5,200 × 3.7 ÷ 1000'),
    ('กำลังไฟฟ้าที่อุปกรณ์ใช้โดยประมาณ', '0.45–0.60 W', 'ESP8266 + GPS + LoRa ขณะทำงานจริง'),
    ('ประสิทธิภาพวงจรจ่ายไฟ', 'ประมาณ 85%', 'เผื่อความสูญเสียจากวงจรจ่ายไฟ สายไฟ และโมดูล'),
    ('พลังงานใช้งานได้หลังหักความสูญเสีย', 'ประมาณ 16.35 Wh', '19.24 × 0.85'),
    ('เวลาใช้งานสูงสุดเชิงทฤษฎี', 'ประมาณ 27–36 ชั่วโมง', '16.35 ÷ 0.60 ถึง 16.35 ÷ 0.45'),
    ('หากใช้งานวันละ 12 ชั่วโมง', 'ประมาณ 2.2–3.0 วัน', 'ไม่คิดพลังงานเสริมจากโซลาร์เซลล์'),
    ('กำลังไฟฟ้าแผงโซลาร์เซลล์', '3.3 W', '6V × 0.55A'),
    ('ประสิทธิภาพรวมของแผงในสภาพจริง', 'ประมาณ 30–50%', 'เผื่อผลจากมุมแดด กระจก ความร้อน เมฆ และเงาบัง'),
    ('กำลังไฟฟ้าจริงโดยประมาณจากแผง', 'ประมาณ 1.0–1.65 W', '3.3W × 30–50%'),
    ('เวลาในการชาร์จ 0–100% เชิงทฤษฎี', 'ประมาณ 12–20 ชั่วโมงแดดจริง', '19.24Wh ÷ 1.65W ถึง 19.24Wh ÷ 1.0W'),
]
table = doc.add_table(rows=len(rows), cols=3)
table.style = 'Table Grid'
table.autofit = False
widths = (Inches(1.75), Inches(2.05), Inches(2.70))
for ri, data in enumerate(rows):
    row = table.rows[ri]
    prevent_row_split(row)
    if ri == 0:
        set_repeat_table_header(row)
    for ci, text in enumerate(data):
        cell = row.cells[ci]
        cell.width = widths[ci]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        if ri == 0:
            shade_cell(cell, 'D9D9D9')
        alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 1 or ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(cell, text, body_template.runs[0], alignment, bold=(ri == 0))
insert_before(anchor, table._tbl)

add_body(doc, anchor, 'พลังงานรวมของแบตเตอรี่สามารถคำนวณได้จากสมการ', body_template)
add_formula_paragraph(doc, anchor, [('E', False), ('battery', True), (' = (Capacity', False), ('mAh', True), (' × V', False), ('battery', True), (') / 1000', False)], body_template)
add_body(doc, anchor, 'แทนค่า', body_template)
add_formula_paragraph(doc, anchor, [('E', False), ('battery', True), (' = (5200 × 3.7) / 1000 = 19.24 Wh', False)], body_template)
add_body(doc, anchor, 'ดังนั้น แบตเตอรี่ขนาด 3.7V 5200mAh มีพลังงานตามสเปกประมาณ 19.24Wh เมื่อนำมาหักความสูญเสียของวงจรจ่ายไฟที่ประสิทธิภาพประมาณ 85% จะได้พลังงานที่สามารถใช้งานได้จริงโดยประมาณดังนี้', body_template)
add_formula_paragraph(doc, anchor, [('E', False), ('usable', True), (' = E', False), ('battery', True), (' × η', False), ('power', True)], body_template)
add_formula_paragraph(doc, anchor, [('E', False), ('usable', True), (' = 19.24 × 0.85 = 16.35 Wh', False)], body_template)
add_body(doc, anchor, 'หากประเมินว่าอุปกรณ์ VIBE ใช้กำลังไฟฟ้าเฉลี่ยประมาณ 0.45–0.60W ระยะเวลาการใช้งานสูงสุดเชิงทฤษฎีสามารถคำนวณได้จาก', body_template)
add_formula_paragraph(doc, anchor, [('Runtime', False), ('theory', True), (' = E', False), ('usable', True), (' / P', False), ('device', True)], body_template)
add_body(doc, anchor, 'กรณีที่อุปกรณ์ใช้กำลังไฟฟ้าประมาณ 0.60W', body_template)
add_formula_paragraph(doc, anchor, [('Runtime', False), ('theory', True), (' = 16.35 / 0.60 = 27.25 ชั่วโมง', False)], body_template)
add_body(doc, anchor, 'กรณีที่อุปกรณ์ใช้กำลังไฟฟ้าประมาณ 0.45W', body_template)
add_formula_paragraph(doc, anchor, [('Runtime', False), ('theory', True), (' = 16.35 / 0.45 = 36.33 ชั่วโมง', False)], body_template)
add_body(doc, anchor, 'ดังนั้น อุปกรณ์ VIBE สามารถทำงานต่อเนื่องได้ประมาณ 27–36 ชั่วโมงในเชิงทฤษฎี หรือประมาณ 2.2–3.0 วัน หากใช้งานวันละ 12 ชั่วโมงโดยไม่คิดพลังงานที่ได้รับเพิ่มเติมจากแผงโซลาร์เซลล์', body_template)
add_body(doc, anchor, 'สำหรับแผงโซลาร์เซลล์ที่ใช้ในระบบต้นแบบ มีสเปก 6V 550mA สามารถคำนวณกำลังไฟฟ้าตามสเปกได้จาก', body_template)
add_formula_paragraph(doc, anchor, [('P', False), ('solar-rated', True), (' = V × I', False)], body_template)
add_formula_paragraph(doc, anchor, [('P', False), ('solar-rated', True), (' = 6 × 0.55 = 3.3 W', False)], body_template)
add_body(doc, anchor, 'อย่างไรก็ตาม ในการติดตั้งใช้งานจริง กำลังไฟฟ้าที่ได้จากแผงโซลาร์เซลล์มักต่ำกว่าค่าตามสเปก เนื่องจากปัจจัยด้านมุมรับแสง เงาบัง ความร้อน สภาพอากาศ และตำแหน่งติดตั้ง เช่น การติดตั้งภายในรถหรือบริเวณที่มีกระจกบังแสง จึงประเมินประสิทธิภาพรวมของแผงในสภาพใช้งานจริงประมาณ 30–50%', body_template)
add_formula_paragraph(doc, anchor, [('P', False), ('solar-real', True), (' = P', False), ('solar-rated', True), (' × K', False), ('solar', True)], body_template)
add_formula_paragraph(doc, anchor, [('P', False), ('solar-real', True), (' = 3.3 × 0.30 ถึง 0.50 ≈ 1.0 ถึง 1.65 W', False)], body_template)
add_body(doc, anchor, 'ดังนั้น แผงโซลาร์เซลล์จะให้กำลังไฟฟ้าจริงโดยประมาณ 1.0–1.65W ในสภาพแสงดีระดับใช้งานจริง และสามารถประเมินเวลาในการชาร์จแบตเตอรี่จาก 0–100% ได้จากสมการ', body_template)
add_formula_paragraph(doc, anchor, [('T', False), ('charge', True), (' = E', False), ('battery', True), (' / P', False), ('solar-real', True)], body_template)
add_body(doc, anchor, 'กรณีที่แผงให้กำลังไฟฟ้าจริงประมาณ 1.65W', body_template)
add_formula_paragraph(doc, anchor, [('T', False), ('charge', True), (' = 19.24 / 1.65 ≈ 11.7 ชั่วโมงแดดจริง', False)], body_template)
add_body(doc, anchor, 'กรณีที่แผงให้กำลังไฟฟ้าจริงประมาณ 1.0W', body_template)
add_formula_paragraph(doc, anchor, [('T', False), ('charge', True), (' = 19.24 / 1.0 ≈ 19.2 ชั่วโมงแดดจริง', False)], body_template)
add_body(doc, anchor, 'ดังนั้น เวลาในการชาร์จแบตเตอรี่จาก 0–100% ด้วยพลังงานแสงอาทิตย์เชิงทฤษฎีจะอยู่ที่ประมาณ 12–20 ชั่วโมงแดดจริง ทั้งนี้ หากอุปกรณ์ยังเปิดใช้งานอยู่ระหว่างการชาร์จ พลังงานบางส่วนจากแผงโซลาร์เซลล์จะถูกใช้เลี้ยงอุปกรณ์ไปพร้อมกัน ทำให้เวลาในการชาร์จจริงอาจนานกว่าค่าที่คำนวณได้', body_template)
add_body(doc, anchor, 'จากการประเมินเชิงทฤษฎี ระบบพลังงานของอุปกรณ์ VIBE มีแนวโน้มรองรับการใช้งานภาคสนามได้ในระดับหนึ่ง โดยแบตเตอรี่สามารถจ่ายพลังงานให้อุปกรณ์ทำงานต่อเนื่องได้ประมาณ 27–36 ชั่วโมง และแผงโซลาร์เซลล์สามารถช่วยเสริมพลังงานหรือชาร์จแบตเตอรี่ได้ตามสภาพแสงจริง อย่างไรก็ตาม ค่าดังกล่าวเป็นเพียงการประเมินจากสเปกอุปกรณ์และสมมติฐานด้านประสิทธิภาพของระบบ จึงควรนำไปเปรียบเทียบกับผลการทดสอบจริง เช่น ระดับแบตเตอรี่ตามเวลา แรงดันไฟฟ้า กราฟการคายประจุ และเวลาในการชาร์จจริง เพื่อประเมินความเหมาะสมของระบบพลังงานสำหรับการติดตั้งบนรถสองแถวในสภาพแวดล้อมจริง', body_template)

# Update only the requested row in the existing technology summary table.
tech_table = next(table for table in doc.tables if table.cell(0, 0).text.strip() == 'หมวดเทคโนโลยี')
power_row = next(row for row in tech_table.rows if row.cells[0].text.strip() == 'Power Management')
set_cell_text(power_row.cells[1], 'Adaptive TX Interval, Sleep Mode, Battery Monitoring, Solar Charging Estimation, Runtime Calculation', body_template.runs[0], WD_ALIGN_PARAGRAPH.LEFT)
set_cell_text(power_row.cells[2], 'ลดการใช้พลังงาน ยืดระยะเวลาการทำงาน ประเมินระยะเวลาการใช้งานของแบตเตอรี่ และประเมินเวลาในการชาร์จด้วยพลังงานแสงอาทิตย์', body_template.runs[0], WD_ALIGN_PARAGRAPH.LEFT)

doc.save(SOURCE)
print(SOURCE)
