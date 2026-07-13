import re
import sys
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def enable_update_fields(document):
    settings = document.settings.element
    node = settings.find(qn('w:updateFields'))
    if node is None:
        node = OxmlElement('w:updateFields')
        settings.append(node)
    node.set(qn('w:val'), 'true')


def configure_heading_style(style):
    style.base_style = document.styles['Normal']
    style.font.name = 'TH Sarabun New'
    style.font.size = None
    style.font.bold = None
    style.paragraph_format.space_before = None
    style.paragraph_format.space_after = None
    style.paragraph_format.keep_with_next = True


def heading_level(text):
    if text in {'กิตติกรรมประกาศ', 'Abstract'}:
        return 1
    if text in {
        'วัตถุประสงค์', 'เป้าหมาย', 'เป้าหมายเชิงระบบ',
        'เป้าหมายเชิงการทดสอบ', 'เป้าหมายด้านการนำไปใช้',
    }:
        return 2
    if re.match(r'^ภาคผนวก\s+[ก-ฮ]', text):
        return 2
    if re.match(r'^[ก-ฮ]\.[0-9]+\s+', text):
        return 3
    if re.match(r'^\d+\.\d+\.\d+(?:\.\d+)?\s+', text):
        return 3
    if re.match(r'^\d+\.\d+\s+', text):
        return 2
    if re.match(r'^\d+\.(?!\d)\s*\S', text):
        return 1
    return None


source, output = sys.argv[1], sys.argv[2]
document = Document(source)
for name in ('Heading 1', 'Heading 2', 'Heading 3'):
    configure_heading_style(document.styles[name])

toc_heading = None
toc_heading_index = None
toc_placeholder = None
first_after_toc = None
count = {1: 0, 2: 0, 3: 0}
for index, paragraph in enumerate(document.paragraphs):
    text = paragraph.text.strip().replace('\n', ' ')
    if text == '3. สารบัญ':
        toc_heading = paragraph
        toc_heading_index = index
        continue
    level = heading_level(text)
    if level is not None:
        paragraph.style = f'Heading {level}'
        count[level] += 1

if toc_heading is None:
    raise RuntimeError('ไม่พบหัวข้อ 3. สารบัญ')

# Reuse the original empty paragraph that follows the TOC heading as the field placeholder.
body = list(document.paragraphs)
for paragraph in body[toc_heading_index + 1:]:
    if not paragraph.text.strip():
        toc_placeholder = paragraph
        continue
    first_after_toc = paragraph
    break

if toc_placeholder is None or first_after_toc is None:
    raise RuntimeError('ไม่พบตำแหน่งว่างสำหรับวางสารบัญ')

toc_placeholder.text = '[[TOC]]'
first_after_toc.paragraph_format.page_break_before = True
enable_update_fields(document)
document.save(output)
print(f'Headings set: H1={count[1]}, H2={count[2]}, H3={count[3]}')
