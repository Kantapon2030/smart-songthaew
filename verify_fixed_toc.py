import sys
from docx import Document
from docx.oxml.ns import qn


document = Document(sys.argv[1])

toc = next(
    table for table in document.tables
    if len(table.columns) == 2 and len(table.rows) == 23
)
grid_widths = [int(col.get(qn('w:w'))) for col in toc._tbl.tblGrid.gridCol_lst]
assert len(grid_widths) == 2
assert grid_widths[1] == 540, grid_widths
assert sum(grid_widths) <= 10000, grid_widths
assert all(width < 10000 for width in grid_widths), grid_widths

page_refs = 0
for row in toc.rows:
    assert row.cells[0].text.strip()
    assert 'PAGEREF toc_entry_' in row.cells[1]._tc.xml
    page_refs += row.cells[1]._tc.xml.count('PAGEREF toc_entry_')
assert page_refs == 23, page_refs
assert 'w:val="nil"' in toc._tbl.xml

toc_heading = next(p for p in document.paragraphs if p.text.strip() == '3. สารบัญ')
main_heading = next(p for p in document.paragraphs if p.text.strip().startswith('4. วัตถุประสงค์และเป้าหมาย'))
children = list(document._body._body)
toc_index = children.index(toc_heading._p)
table_index = children.index(toc._tbl)
main_index = children.index(main_heading._p)
assert toc_index < table_index < main_index, (toc_index, table_index, main_index)
section_break = children[table_index + 1]
assert section_break.tag == qn('w:p')
p_pr = section_break.find(qn('w:pPr'))
assert p_pr is not None and p_pr.find(qn('w:sectPr')) is not None

settings = document.settings.element.xml
assert 'w:updateFields' in settings
assert len(document.sections) == 3
print('PASS: compact TOC fits print width, has 23 page references, is borderless, and remains before the main-content section break.')
